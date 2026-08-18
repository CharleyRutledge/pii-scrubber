"""Plain-text extraction for supported document types."""

import csv
import json
from pathlib import Path

_TEXT_SUFFIXES = {".txt", ".md"}


def extract_text(path: str | Path) -> str:
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix in _TEXT_SUFFIXES:
        return path.read_text(encoding="utf-8")
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".csv":
        return _extract_csv(path)
    if suffix == ".json":
        return _extract_json(path)

    raise ValueError(f"Unsupported file type: {suffix}")


def _extract_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError(
            "PDF support requires pdfplumber. Install with: pip install pii-scrubber[pdf]"
        ) from exc

    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _extract_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError(
            "DOCX support requires python-docx. Install with: pip install pii-scrubber[docx]"
        ) from exc

    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def _extract_csv(path: Path) -> str:
    with path.open(newline="", encoding="utf-8") as f:
        reader = csv.reader(f)
        return "\n".join(", ".join(row) for row in reader)


def _extract_json(path: Path) -> str:
    data = json.loads(path.read_text(encoding="utf-8"))
    return _flatten_json(data)


def _flatten_json(data) -> str:
    """Flatten arbitrary JSON into newline-separated text so NER/regex can scan it."""
    parts: list[str] = []

    def walk(value):
        if isinstance(value, dict):
            for v in value.values():
                walk(v)
        elif isinstance(value, list):
            for v in value:
                walk(v)
        elif value is not None:
            parts.append(str(value))

    walk(data)
    return "\n".join(parts)
