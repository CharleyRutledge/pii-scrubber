"""Format-preserving redaction: write a redacted copy of a document in its
original file format, instead of flattened plain text.

.txt/.md  -> redacted text written back out
.docx     -> paragraph/table-cell text replaced in place, structure kept
.csv      -> redacted per-cell, rows/columns kept
.json     -> string values redacted in place, structure/keys kept
.pdf      -> original layout/images kept; black boxes burned over PII text
.html/.htm -> tags/attributes kept, only visible text nodes redacted
"""

import csv
import json
import warnings
from pathlib import Path

from .core import scrub
from .html_tools import split_html_segments
from .open_file import open_with_default_app


def redact_file(
    path: str | Path,
    output_path: str | Path | None = None,
    use_ner: bool = True,
    ner_model: str = "en_core_web_sm",
    ner_labels: set[str] | None = None,
    ocr: bool = False,
    open_after: bool = False,
) -> Path:
    """Write a redacted copy of `path` in the same file format and return its path.

    If `output_path` is omitted, writes next to the original as
    "<stem>_redacted<suffix>".

    `ocr` (PDF only) additionally OCRs embedded images and fully blacks out
    any image whose recognized text contains PII (e.g. a photographed ID or
    a screenshot). Requires the `ocr` extra and a system Tesseract install -
    see `pii_scrubber.ocr` for setup. Off by default since it's slower and
    pulls in an extra system dependency.

    `open_after` launches the redacted file in whatever app the OS has
    associated with it once writing finishes, so you can immediately
    eyeball the result. Off by default - opt in explicitly, since silently
    launching an app is surprising behavior for anything running
    unattended (scripts, CI, batch jobs).
    """
    path = Path(path)
    suffix = path.suffix.lower()
    if output_path is None:
        output_path = path.with_name(f"{path.stem}_redacted{path.suffix}")
    output_path = Path(output_path)

    def _scrub(text: str) -> str:
        return scrub(text, use_ner=use_ner, ner_model=ner_model, ner_labels=ner_labels).text

    if suffix in {".txt", ".md"}:
        _redact_plain_text(path, output_path, _scrub)
    elif suffix in {".html", ".htm"}:
        _redact_html(path, output_path, _scrub)
    elif suffix == ".docx":
        _redact_docx(path, output_path, _scrub)
    elif suffix == ".csv":
        _redact_csv(path, output_path, _scrub)
    elif suffix == ".json":
        _redact_json(path, output_path, _scrub)
    elif suffix == ".pdf":
        _redact_pdf(path, output_path, use_ner, ner_model, ner_labels, ocr)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    if open_after:
        open_with_default_app(output_path)

    return output_path


def _redact_plain_text(path: Path, output_path: Path, scrub_fn) -> None:
    text = path.read_text(encoding="utf-8")
    output_path.write_text(scrub_fn(text), encoding="utf-8")


def _redact_html(path: Path, output_path: Path, scrub_fn) -> None:
    html = path.read_text(encoding="utf-8")
    segments = split_html_segments(html)
    rebuilt = "".join(
        seg if is_tag_or_skipped else scrub_fn(seg) for is_tag_or_skipped, seg in segments
    )
    output_path.write_text(rebuilt, encoding="utf-8")


def _set_paragraph_text(paragraph, new_text: str) -> None:
    """Replace a python-docx paragraph's visible text with `new_text`,
    collapsing all runs into one (per-run formatting is not preserved
    since a PII span can straddle multiple runs).
    """
    if not paragraph.runs:
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _redact_docx(path: Path, output_path: Path, scrub_fn) -> None:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError(
            "DOCX support requires python-docx. Install with: pip install pii-scrubber[docx]"
        ) from exc

    document = docx.Document(str(path))

    for paragraph in document.paragraphs:
        if paragraph.text:
            _set_paragraph_text(paragraph, scrub_fn(paragraph.text))

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        _set_paragraph_text(paragraph, scrub_fn(paragraph.text))

    document.save(str(output_path))


def _redact_csv(path: Path, output_path: Path, scrub_fn) -> None:
    with path.open(newline="", encoding="utf-8") as f:
        rows = list(csv.reader(f))

    redacted_rows = [[scrub_fn(cell) for cell in row] for row in rows]

    with output_path.open("w", newline="", encoding="utf-8") as f:
        csv.writer(f).writerows(redacted_rows)


def _redact_json(path: Path, output_path: Path, scrub_fn) -> None:
    data = json.loads(path.read_text(encoding="utf-8"))

    def walk(value):
        if isinstance(value, dict):
            return {k: walk(v) for k, v in value.items()}
        if isinstance(value, list):
            return [walk(v) for v in value]
        if isinstance(value, str):
            return scrub_fn(value)
        return value

    redacted = walk(data)
    output_path.write_text(json.dumps(redacted, indent=2), encoding="utf-8")


def _redact_pdf(
    path: Path,
    output_path: Path,
    use_ner: bool,
    ner_model: str,
    ner_labels: set[str] | None,
    ocr: bool,
) -> None:
    try:
        import pymupdf as fitz
    except ImportError as exc:
        raise RuntimeError(
            "PDF support requires PyMuPDF. Install with: pip install pii-scrubber[pdf]"
        ) from exc

    with fitz.open(str(path)) as doc:
        for page_number, page in enumerate(doc):
            page_text = page.get_text()
            has_broken_glyphs = "�" in page_text

            if page_text.strip():
                result = scrub(
                    page_text, use_ner=use_ner, ner_model=ner_model, ner_labels=ner_labels
                )
                pii_strings = {e.text for e in result.entities if e.text.strip()}

                for pii_text in pii_strings:
                    for rect in page.search_for(pii_text):
                        page.add_redact_annot(rect, fill=(0, 0, 0))

            if has_broken_glyphs:
                if ocr:
                    _redact_page_via_ocr(
                        page, fitz, use_ner, ner_model, ner_labels
                    )
                else:
                    warnings.warn(
                        f"Page {page_number + 1} has text the PDF's font can't decode "
                        "properly (shows up as �) - PII inside it may not be fully "
                        "redacted. Pass ocr=True (requires the `ocr` extra + Tesseract) "
                        "to redact via OCR instead.",
                        stacklevel=2,
                    )

            if ocr:
                _redact_page_images(doc, page, fitz, use_ner, ner_model, ner_labels)

            page.apply_redactions()

        doc.save(str(output_path))


def _redact_page_via_ocr(page, fitz, use_ner, ner_model, ner_labels) -> None:
    """Fallback for pages whose embedded font can't be decoded cleanly
    (native text extraction yields U+FFFD replacement chars). Rasterizes
    the page and redacts via OCR word bounding boxes instead, so PII isn't
    silently missed just because the text layer is broken.
    """
    from .ocr import ocr_words_with_boxes

    zoom = 3  # ~300dpi equivalent for reliable OCR accuracy
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
    img_bytes = pix.tobytes("png")

    try:
        words = ocr_words_with_boxes(img_bytes)
    except Exception:
        return
    if not words:
        return

    # Reconstruct OCR'd text with a char-offset -> word-index map so PII
    # spans (which can cover multiple words) can be traced back to boxes.
    # Start a new line in the reconstruction whenever OCR's own line/block
    # numbering changes, instead of flattening everything to one line -
    # otherwise line-anchored regexes (e.g. address detection) can't tell
    # where a line ends and could match across the whole page.
    parts = []
    offsets = []  # (start, end, word_index)
    cursor = 0
    prev_line_key = None
    for i, w in enumerate(words):
        line_key = (w["block_num"], w["line_num"])
        if prev_line_key is not None and line_key != prev_line_key:
            parts.append("\n")
            cursor += 1
        elif prev_line_key is not None:
            parts.append(" ")
            cursor += 1
        prev_line_key = line_key

        start = cursor
        parts.append(w["text"])
        cursor += len(w["text"])
        offsets.append((start, cursor, i))
    reconstructed = "".join(parts)

    result = scrub(reconstructed, use_ner=use_ner, ner_model=ner_model, ner_labels=ner_labels)

    for entity in result.entities:
        matching_words = [
            words[i] for start, end, i in offsets if start < entity.end and end > entity.start
        ]
        if not matching_words:
            continue

        # One rect per OCR line the match touches, rather than a single box
        # spanning every matched word - keeps a mis-scoped match from ever
        # covering unrelated content between its first and last line.
        by_line: dict[tuple, list[dict]] = {}
        for w in matching_words:
            by_line.setdefault((w["block_num"], w["line_num"]), []).append(w)

        for line_words in by_line.values():
            x0 = min(w["left"] for w in line_words) / zoom
            y0 = min(w["top"] for w in line_words) / zoom
            x1 = max(w["left"] + w["width"] for w in line_words) / zoom
            y1 = max(w["top"] + w["height"] for w in line_words) / zoom
            page.add_redact_annot(fitz.Rect(x0, y0, x1, y1), fill=(0, 0, 0))


def _redact_page_images(doc, page, fitz, use_ner, ner_model, ner_labels) -> None:
    """OCR each image on the page; if any PII is found in it, black out the
    whole image region. We can't reliably map OCR'd text back to pixel
    coordinates within the image, so this over-redacts (whole image, not
    just the PII substring) rather than risk leaving PII visible.
    """
    from .ocr import ocr_image_bytes

    for img in page.get_images(full=True):
        xref = img[0]
        try:
            image_bytes = doc.extract_image(xref)["image"]
            ocr_text = ocr_image_bytes(image_bytes)
        except Exception:
            continue

        if not ocr_text.strip():
            continue

        result = scrub(ocr_text, use_ner=use_ner, ner_model=ner_model, ner_labels=ner_labels)
        if not result.entities:
            continue

        for rect in page.get_image_rects(xref):
            page.add_redact_annot(rect, fill=(0, 0, 0))
