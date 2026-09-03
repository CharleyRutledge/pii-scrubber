import json

import docx
import pytest

from pii_scrubber import redact_file


def test_redact_txt(tmp_path):
    src = tmp_path / "doc.txt"
    src.write_text("Email jane.doe@example.com now.", encoding="utf-8")

    out = redact_file(src, use_ner=False)

    assert out == tmp_path / "doc_redacted.txt"
    assert "[EMAIL]" in out.read_text(encoding="utf-8")
    assert "jane.doe@example.com" not in out.read_text(encoding="utf-8")


def test_redact_txt_custom_output(tmp_path):
    src = tmp_path / "doc.txt"
    src.write_text("Call 555-123-4567.", encoding="utf-8")
    dest = tmp_path / "custom.txt"

    out = redact_file(src, output_path=dest, use_ner=False)

    assert out == dest
    assert "[PHONE]" in dest.read_text(encoding="utf-8")


def test_redact_csv_preserves_structure(tmp_path):
    src = tmp_path / "doc.csv"
    src.write_text("name,email\nJane,jane@example.com\n", encoding="utf-8")

    out = redact_file(src, use_ner=False)
    content = out.read_text(encoding="utf-8")

    assert "name,email" in content
    assert "jane@example.com" not in content
    assert "[EMAIL]" in content


def test_redact_json_preserves_keys(tmp_path):
    src = tmp_path / "doc.json"
    src.write_text(
        json.dumps({"user": {"email": "jane@example.com", "age": 30}}),
        encoding="utf-8",
    )

    out = redact_file(src, use_ner=False)
    data = json.loads(out.read_text(encoding="utf-8"))

    assert data["user"]["age"] == 30
    assert data["user"]["email"] == "[EMAIL]"


def test_redact_docx_preserves_paragraph_structure(tmp_path):
    src = tmp_path / "doc.docx"
    document = docx.Document()
    document.add_paragraph("Contact jane.doe@example.com for details.")
    document.add_paragraph("This paragraph has no PII.")
    document.save(str(src))

    out = redact_file(src, use_ner=False)
    result = docx.Document(str(out))

    assert len(result.paragraphs) == 2
    assert "[EMAIL]" in result.paragraphs[0].text
    assert "jane.doe@example.com" not in result.paragraphs[0].text
    assert result.paragraphs[1].text == "This paragraph has no PII."


def test_redact_html_keeps_tags_and_redacts_only_text(tmp_path):
    src = tmp_path / "page.html"
    src.write_text(
        "<html><body>\n"
        "<h1>Contact</h1>\n"
        '<p class="lead">Reach jane@example.com or 555-123-4567.</p>\n'
        "</body></html>\n",
        encoding="utf-8",
    )

    out = redact_file(src, use_ner=False)
    content = out.read_text(encoding="utf-8")

    # Tags/attributes untouched...
    assert "<h1>Contact</h1>" in content
    assert 'class="lead"' in content
    # ...but the PII in the visible text is gone.
    assert "jane@example.com" not in content
    assert "555-123-4567" not in content
    assert "[EMAIL]" in content
    assert "[PHONE]" in content


def test_redact_docx_redacts_table_cells(tmp_path):
    src = tmp_path / "table.docx"
    document = docx.Document()
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Jane Doe"
    table.rows[0].cells[1].text = "jane@example.com"
    document.save(str(src))

    out = redact_file(src, use_ner=False)
    result = docx.Document(str(out))
    cell_texts = [c.text for row in result.tables[0].rows for c in row.cells]

    assert "jane@example.com" not in cell_texts
    assert any("[EMAIL]" in t for t in cell_texts)


def test_redact_pdf_removes_pii_from_text_layer(tmp_path):
    pymupdf = pytest.importorskip("pymupdf")
    src = tmp_path / "doc.pdf"
    pdf = pymupdf.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Email jane@example.com or call 555-123-4567.")
    pdf.save(str(src))
    pdf.close()

    out = redact_file(src, use_ner=False)

    with pymupdf.open(str(out)) as result:
        text = "".join(p.get_text() for p in result)
    assert "jane@example.com" not in text
    assert "555-123-4567" not in text


def test_redact_unsupported_extension_raises(tmp_path):
    src = tmp_path / "data.xyz"
    src.write_text("Email jane@example.com", encoding="utf-8")

    with pytest.raises(ValueError, match="Unsupported file type"):
        redact_file(src, use_ner=False)


def test_redact_open_after_launches_the_output_file(tmp_path, monkeypatch):
    # Mocked so this never actually launches a real application.
    opened = []
    monkeypatch.setattr("pii_scrubber.redact.open_with_default_app", opened.append)

    src = tmp_path / "doc.txt"
    src.write_text("Email jane.doe@example.com now.", encoding="utf-8")

    out = redact_file(src, use_ner=False, open_after=True)

    assert opened == [out]


def test_redact_does_not_open_by_default(tmp_path, monkeypatch):
    opened = []
    monkeypatch.setattr("pii_scrubber.redact.open_with_default_app", opened.append)

    src = tmp_path / "doc.txt"
    src.write_text("Email jane.doe@example.com now.", encoding="utf-8")

    redact_file(src, use_ner=False)

    assert opened == []
