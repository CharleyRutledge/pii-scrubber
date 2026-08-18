"""Runs the CLI end-to-end against real fixture files, captures the actual
terminal output of each step, renders it as an animated GIF of a terminal
session, and writes a step-by-step Markdown guide from the same captured
data - so the GIF and the docs can never drift out of sync with each other
or with what the CLI actually does.

Usage: python tools/record_demo.py [output_dir]
"""

import json
import subprocess
import sys
import tempfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont

# --- terminal look ---
_BG = (12, 12, 16)
_FG = (223, 223, 223)
_PROMPT = (98, 209, 150)
_DIM = (140, 140, 150)
_FONT_SIZE = 18
_LINE_HEIGHT = 24
_PADDING = 18
_WIDTH = 900

_FONT_CANDIDATES = [
    r"C:\Windows\Fonts\consola.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationMono-Regular.ttf",
    "/System/Library/Fonts/Menlo.ttc",
]


def _load_font() -> ImageFont.FreeTypeFont:
    for path in _FONT_CANDIDATES:
        if Path(path).is_file():
            return ImageFont.truetype(path, _FONT_SIZE)
    return ImageFont.load_default(size=_FONT_SIZE)


def _run(command: list[str], cwd: Path) -> tuple[str, int]:
    proc = subprocess.run(
        command, cwd=cwd, capture_output=True, text=True, encoding="utf-8"
    )
    combined = proc.stdout + proc.stderr
    return combined, proc.returncode


_INTAKE_TEXT = (
    "Patient Intake Form\n\n"
    "Name: Jane Doe\n"
    "Email: jane.doe@example.com\n"
    "Phone: 555-123-4567\n"
    "SSN: 123-45-6789\n"
    "Employer: Acme Corp, located in Seattle.\n"
)


def _setup_fixtures(workspace: Path) -> None:
    (workspace / "patient_intake.txt").write_text(_INTAKE_TEXT, encoding="utf-8")

    (workspace / "notes.md").write_text(
        "# Follow-up notes\n\n"
        "Reach **Jane Doe** at jane.doe@example.com or 555-123-4567 "
        "before the Friday deadline.\n",
        encoding="utf-8",
    )

    (workspace / "profile.html").write_text(
        "<html><body>\n"
        "<h1>Contact card</h1>\n"
        "<p>Jane Doe - jane.doe@example.com - 555-123-4567</p>\n"
        "<p>LinkedIn: linkedin.com/in/janedoe</p>\n"
        "</body></html>\n",
        encoding="utf-8",
    )

    (workspace / "contacts.csv").write_text(
        "name,email,phone\n"
        "Jane Doe,jane.doe@example.com,555-123-4567\n"
        "John Smith,john.smith@example.com,555-987-6543\n",
        encoding="utf-8",
    )

    (workspace / "record.json").write_text(
        json.dumps(
            {
                "patient": {
                    "name": "Jane Doe",
                    "email": "jane.doe@example.com",
                    "phone": "555-123-4567",
                }
            },
            indent=2,
        ),
        encoding="utf-8",
    )

    import docx

    document = docx.Document()
    document.add_paragraph("Employment Letter")
    document.add_paragraph(
        "This confirms Jane Doe (jane.doe@example.com, 555-123-4567) "
        "is employed at Acme Corp."
    )
    document.save(str(workspace / "letter.docx"))

    import pymupdf

    pdf = pymupdf.open()
    page = pdf.new_page()
    page.insert_text(
        (72, 72),
        "Please contact Jane Doe at jane.doe@example.com or 555-123-4567.",
    )
    pdf.save(str(workspace / "notice.pdf"))
    pdf.close()


class _Step:
    def __init__(self, description: str, command: list[str], display_command: str):
        self.description = description
        self.command = command
        self.display_command = display_command
        self.output = ""
        self.returncode = 0


_SCRUB_DEMO_FILES = [
    ("patient_intake.txt", "plain text"),
    ("notes.md", "Markdown"),
    ("profile.html", "HTML"),
    ("contacts.csv", "CSV"),
    ("record.json", "JSON"),
    ("letter.docx", "Word (.docx)"),
    ("notice.pdf", "PDF"),
]


def _build_steps(python_exe: str) -> list[_Step]:
    cli = [python_exe, "-m", "pii_scrubber.cli"]
    steps = [
        _Step(
            "See what the CLI offers.",
            [*cli, "--help"],
            "pii-scrubber --help",
        ),
    ]

    for filename, format_label in _SCRUB_DEMO_FILES:
        steps.append(
            _Step(
                f"Scrub a {format_label} file and preview the redacted text + "
                "a summary count, without writing any file.",
                [*cli, "scrub", filename],
                f"pii-scrubber scrub {filename}",
            )
        )

    steps.append(
        _Step(
            "Write a redacted copy of the file, format preserved.",
            [*cli, "redact", "patient_intake.txt"],
            "pii-scrubber redact patient_intake.txt",
        )
    )
    steps.append(
        _Step(
            "Confirm the original PII is gone from the redacted file.",
            [python_exe, "-c",
             "print(open('patient_intake_redacted.txt', encoding='utf-8').read())"],
            "cat patient_intake_redacted.txt",
        )
    )
    return steps


def _wrap_text(text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines = []
    for raw_line in text.splitlines() or [""]:
        if font.getlength(raw_line) <= max_width:
            lines.append(raw_line)
            continue
        current = ""
        for word in raw_line.split(" "):
            candidate = f"{current} {word}".strip()
            if font.getlength(candidate) <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        lines.append(current)
    return lines


def _build_line_snapshots(
    steps: list[_Step], font: ImageFont.FreeTypeFont
) -> list[tuple[list[tuple[str, tuple[int, int, int]]], int]]:
    """Returns (accumulated_lines, hold_frames) for each snapshot point, in order."""
    max_text_width = _WIDTH - 2 * _PADDING
    all_lines: list[tuple[str, tuple[int, int, int]]] = []
    snapshots: list[tuple[list[tuple[str, tuple[int, int, int]]], int]] = []

    for step in steps:
        all_lines.append(("", _FG))
        for wrapped in _wrap_text(f"# {step.description}", font, max_text_width):
            all_lines.append((wrapped, _DIM))
        all_lines.append((f"$ {step.display_command}", _PROMPT))
        snapshots.append((list(all_lines), 8))

        for wrapped in _wrap_text(step.output.rstrip("\n"), font, max_text_width):
            all_lines.append((wrapped, _FG))
        snapshots.append((list(all_lines), 14))

    return snapshots


def _render_frames(steps: list[_Step], font: ImageFont.FreeTypeFont) -> list[Image.Image]:
    snapshots = _build_line_snapshots(steps, font)
    max_lines = max(len(lines) for lines, _ in snapshots)
    canvas_size = (_WIDTH, max(_PADDING * 2 + max_lines * _LINE_HEIGHT, 200))

    frames: list[Image.Image] = []
    for lines, hold_frames in snapshots:
        img = Image.new("RGB", canvas_size, _BG)
        draw = ImageDraw.Draw(img)
        y = _PADDING
        for line_text, color in lines:
            draw.text((_PADDING, y), line_text, font=font, fill=color)
            y += _LINE_HEIGHT
        for _ in range(hold_frames):
            frames.append(img)

    return frames


def generate(docs_dir: Path) -> dict:
    """Runs the demo end-to-end, writes assets/demo.gif and
    USAGE_WALKTHROUGH.md into docs_dir, and returns the captured steps for
    test assertions.
    """
    docs_dir.mkdir(parents=True, exist_ok=True)
    assets_dir = docs_dir / "assets"
    assets_dir.mkdir(exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="pii-scrubber-demo-") as tmp:
        workspace = Path(tmp)
        _setup_fixtures(workspace)

        steps = _build_steps(sys.executable)
        for step in steps:
            step.output, step.returncode = _run(step.command, cwd=workspace)

    font = _load_font()
    frames = _render_frames(steps, font)
    gif_path = assets_dir / "demo.gif"
    frames[0].save(
        gif_path,
        save_all=True,
        append_images=frames[1:],
        duration=120,
        loop=0,
    )

    doc_path = docs_dir / "USAGE_WALKTHROUGH.md"
    doc_path.write_text(_render_markdown(steps), encoding="utf-8")

    return {"steps": steps, "gif_path": gif_path, "doc_path": doc_path}


def _render_markdown(steps: list[_Step]) -> str:
    parts = [
        "# pii-scrubber: step-by-step walkthrough",
        "",
        "This walkthrough is generated directly from a real, passing run of the "
        "CLI (see `tools/record_demo.py`) - every command and output below is "
        "the actual output produced by that run, not hand-written.",
        "",
        "![demo](assets/demo.gif)",
        "",
    ]
    for i, step in enumerate(steps, start=1):
        parts.append(f"## {i}. {step.description}")
        parts.append("")
        parts.append("```console")
        parts.append(f"$ {step.display_command}")
        parts.append(step.output.rstrip("\n"))
        parts.append("```")
        parts.append("")
    return "\n".join(parts)


if __name__ == "__main__":
    target = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("docs")
    result = generate(target)
    print(f"Wrote {result['gif_path']} and {result['doc_path']}")
